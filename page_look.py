import aspose.words as aw
import docx
import math as m
from copy import deepcopy
from docx.shared import Mm
from docx.shared import Length,Pt
from docx.enum.text import WD_LINE_SPACING

def csv_to_doxc(file_rows,open_file_name,vl_name,vl_voltage,vl_uchastok,vl_podstation,name_town,type_looking,date,name_worker1,name_worker2,name_master,name_ingeener):
    open_file_name=open_file_name[:-5]
    def make_page(lst_rows,file_name,open_file_name,vl_name,vl_voltage,vl_uchastok,vl_podstation,name_town,type_looking,date,name_worker1,name_worker2,name_master,name_ingeener):
        print('Create a new Word document.')
        print(lst_rows)
        doc = aw.Document()

        # Create document builder.
        builder = aw.DocumentBuilder(doc)


        # Создание страницы

        # Шапка страницы
        builder.write(
            "Филиал БЭС                                                                                                      Сельский РЭС\n\n")
        builder.write("                                                   ЛИСТОК ОСМОТРА ВЛ 0,4-10 кВ\n")
        builder.write(f"ВЛ {vl_voltage} кВ № {vl_name} участок {vl_uchastok}\n")
        builder.write(f"От подстанции {vl_podstation} н.п. {name_town}\n")
        builder.write(f"Вид осмотра {type_looking} Дата осмотра {date}\n\n")

        # Таблица страницы
        table = builder.start_table()
        # Insert cell.
        builder.insert_cell()
        # Table wide formatting must be applied after at least one row is present in the table.
        table.left_indent = 1.0
        # Set height and define the height rule for the header row.
        builder.row_format.height = 10.0
        builder.row_format.height_rule = aw.HeightRule.AT_LEAST
        # Set alignment and font settings.
        builder.paragraph_format.alignment = aw.ParagraphAlignment.CENTER
        builder.font.size = 8.0
        builder.font.name = "Arial"
        builder.font.bold = True

        builder.cell_format.width = 35.0
        builder.write("Номера опор,\nпролетов")

        # We don't need to specify this cell's width because it's inherited from the previous cell.
        builder.font.size = 10
        builder.insert_cell()
        builder.cell_format.width = 250.0
        builder.write("Выявленные дефекты")

        builder.insert_cell()
        builder.cell_format.width = 120.0
        builder.write("Состояние трассы")
        builder.end_row()
        for i in lst_rows:
            builder.cell_format.width = 35.0
            builder.cell_format.vertical_alignment = aw.tables.CellVerticalAlignment.CENTER

            # Reset height and define a different height rule for table body.
            builder.row_format.height = 10.0
            builder.row_format.height_rule = aw.HeightRule.AUTO
            builder.insert_cell()

            # Reset font formatting.
            builder.font.size = 10
            builder.font.bold = False

            builder.write(i[0])

            builder.insert_cell()
            builder.cell_format.width = 250.0
            builder.write(i[1])

            builder.insert_cell()
            builder.cell_format.width = 120.0
            builder.write(i[2])
            builder.end_row()


        # End table.
        builder.end_table()
        # Концовка страницы
        builder.write("Осмотр произвели:\n")
        builder.write(f"ФИО {name_worker1} Подпись _____________\n")
        builder.write(f"ФИО {name_worker2} Подпись _____________\n")
        builder.write(f"Листок осмотра принял ___________ {name_master} ________________\n")
        builder.write("Результаты осмотра ппроанализированы ________________________\n")
        builder.write(f"Главный инженер РЭС {name_ingeener} ________________\n")

        # Save the document.
        doc.save(f'{open_file_name}, часть{name_file}.docx')
        template  = docx.Document(f'{open_file_name}, часть{name_file}.docx')
        doc1=docx.Document()



        paragraph = doc1.add_paragraph()



        tbl = template.tables[0]._tbl
        new_tabl = deepcopy(tbl)

        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_fmt = paragraph.paragraph_format
        p_fmt.line_spacing = Pt(1)
        isinstance(p_fmt.line_spacing, Length)
        p_fmt.line_spacing = 1

        fmt = paragraph.paragraph_format
        fmt.first_line_indent = Mm(1)
        fmt.space_before = Mm(1)
        fmt.space_after = Mm(1)


        paragraph._p.addnext(new_tabl)

        doc1.save(f'test {open_file_name}, часть{name_file}.docx')
    lst_rows=[]
    count=0
    name_file=1
    for i in file_rows:
        if count+m.ceil(len(i[1])/55)+1<39:

            fin_wrt=''
            for j in i[1]:
                if len(fin_wrt)%55==0 and len(fin_wrt)>=55:
                    fin_wrt+=('\n'+j)
                else:
                    fin_wrt+=j
            lst_rows.append([i[0],fin_wrt,i[2]])
            count += m.ceil(len(i[1])/55)+1

        else:
            print('отправлено в печать')
            make_page(lst_rows,name_file,open_file_name,vl_name,vl_voltage,vl_uchastok,vl_podstation,name_town,type_looking,date,name_worker1,name_worker2,name_master,name_ingeener)
            name_file+=1
            fin_wrt = ''
            for j in i[1]:
                if len(fin_wrt) % 55 == 0 and len(fin_wrt) >= 55:
                    fin_wrt += ('\n' + j)
                else:
                    fin_wrt += j
            lst_rows=[[i[0],fin_wrt,i[2]]]
            count=m.ceil(len(i[1])/55)+1
    if len(lst_rows)>0:
        print(lst_rows)
        print('отправлено в печать,конец файла')
        make_page(lst_rows,name_file, open_file_name,vl_name,vl_voltage,vl_uchastok,vl_podstation,name_town,type_looking,date,name_worker1,name_worker2,name_master,name_ingeener)


#csv_to_doxc([[str(i),'m'*120,' '] for i in range(1,19)],'пробный','969','10','1-70','Козлякевичи','Козлякевичи','Плановый','28.01.2023','Кемза А.Г',' ','Довгер С.А.','Говдей А.Н')
